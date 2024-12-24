import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Загрузка Swagger JSON
with open('swagger.json', 'r', encoding='utf-8') as f:
    swagger_data = json.load(f)

document = Document()

# Вспомогательные функции
def set_cell_background(cell, color):
    """Установить фон для ячейки."""
    cell_properties = cell._element.tcPr
    cell_shading = OxmlElement('w:shd')
    cell_shading.set(qn('w:fill'), color)
    cell_properties.append(cell_shading)

def set_cell_font(cell, bold=False, color=None, size=10):
    """Установить шрифт в ячейке."""
    run = cell.paragraphs[0].runs[0]
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.size = Pt(size)

def set_paragraph_style(paragraph, color=None, bold=False):
    """Стилизация параграфа."""
    run = paragraph.runs[0]
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def create_table_with_header(document, headers, widths):
    """Создать таблицу с заголовками."""
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_background(cell, 'BDD6EE')  # Светло-голубой фон
        set_cell_font(cell, bold=True, size=10)
        cell.width = Inches(widths[idx])
    return table

# Генерация документации
document.add_heading('API Documentation', level=1)

# Обработка Paths
paths = swagger_data.get('paths', {})
for path, methods in paths.items():
    for method, details in methods.items():
        # Заголовок метода
        document.add_heading(f"{method.upper()} {path}", level=2)

        # Описание
        summary = details.get('summary', '')
        description = details.get('description', '')
        document.add_paragraph(summary, style='Heading 3')
        if description:
            document.add_paragraph(description)

        # Таблица параметров
        if 'parameters' in details:
            document.add_heading('Parameters', level=3)
            table = create_table_with_header(
                document,
                ['Name', 'In', 'Type', 'Description', 'Required'],
                [1.5, 1, 1, 3, 1]
            )
            for param in details['parameters']:
                row_cells = table.add_row().cells
                row_cells[0].text = param.get('name', '-')
                row_cells[1].text = param.get('in', '-')
                row_cells[2].text = param.get('type', param.get('schema', {}).get('type', '-'))
                row_cells[3].text = param.get('description', '-')
                row_cells[4].text = 'Yes' if param.get('required') else 'No'

        # Таблица ответов
        if 'responses' in details:
            document.add_heading('Responses', level=3)
            table = create_table_with_header(
                document,
                ['HTTP Code', 'Description', 'Schema'],
                [1.5, 3, 3]
            )
            for code, response in details['responses'].items():
                row_cells = table.add_row().cells
                row_cells[0].text = code
                row_cells[1].text = response.get('description', '-')
                schema = response.get('content', {}).get('application/json', {}).get('schema', {})
                row_cells[2].text = schema.get('$ref', schema.get('type', '-'))

# Добавление моделей (Definitions)
definitions = swagger_data.get('definitions', swagger_data.get('components', {}).get('schemas', {}))
if definitions:
    document.add_heading('Models', level=1)
    for model_name, model_details in definitions.items():
        document.add_heading(model_name, level=2)

        table = create_table_with_header(
            document,
            ['Property', 'Type', 'Description'],
            [2, 2, 4]
        )

        properties = model_details.get('properties', {})
        for prop_name, prop_details in properties.items():
            row_cells = table.add_row().cells
            row_cells[0].text = prop_name
            row_cells[1].text = prop_details.get('type', '-')
            row_cells[2].text = prop_details.get('description', '-')

# Сохранение документа
document.save('swagger_updated.docx')

